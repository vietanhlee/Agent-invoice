from langchain_core.tools import tool
from docx import Document
from dotenv import load_dotenv
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
import pandas as pd
from typing import List, Optional
from pydantic import BaseModel, Field
from langgraph.graph import state
from langchain_google_genai import GoogleGenerativeAI
@tool
def make_file_txt(filename: str, content: str) -> str:
    """ Tạo một file định dạng txt

    Args:
        filename (str): tên file để tạo và ghi nội dung vào đó khong cần phần mở rộng
        content (str): nội dung muốn ghi vào file 
    """
    try:
        with open(filename + '.txt', 'w') as f:
            f.write(content)
    except Exception as e:
        return f'Lỗi {e} khi tạo file {filename}'
    else:
        return f'Đã tạo thành công file {filename}'
          
@tool
def make_file_docx(filename: str, content: str) -> None:
    """ Tạo một file định dạng docx hoặc doc 

    Args:
        filename (str): tên file để tạo và ghi nội dung vào đó không cân phần mở rộng
        content (str): nội dung muốn ghi vào file 
    """
    try:
        document = Document()
        document.add_paragraph(text= content)
        document.save(filename + '.docx')
    except Exception as e:
        return f'Lỗi {e} khi tạo file {filename}'
    else:
        return f'Đã tạo thành công file {filename}'


def read_excel_data(file_path: str) -> List[List[str]]:
    """Đọc nội dung từ file Excel hoá đơn và trả về nội dung dưới dạng list, list chứa các chuỗi
    Args:
        file_name (str): Tên file Excel cần đọc

    Returns:
        list[list[str]]: Nội dung của file Excel theo định dạng list, mỗi list con là một hàng trong file Excel
        Nếu có lỗi xảy ra, trả về thông báo lỗi
    """
    print(f"--- TOOL: Đang đọc file: {file_path} ---")
    try:
        df = pd.read_excel(file_path)
        # Chuyển DataFrame thành chuỗi để LLM có thể đọc dễ dàng
        return df.values.tolist()
    except FileNotFoundError:
        return "Lỗi: Không tìm thấy file Excel."
    except Exception as e:
        return f"Lỗi khi đọc file Excel: {e}"

class InvoiceDetails(BaseModel):
    """Cấu trúc dữ liệu chi tiết của một hoá đơn."""
    du_lieu_bang: Optional[List[List[str]]] = Field(default= [["SP", "SL", "Đơn giá", "TT"]], description= 'Dữ liệu hiển thị cho bảng là một list các list, mỗi list con chứa thông tin về sản phẩm định dạng string\
            tiêu đề bảng thường là: STT, Tên hàng hoá, dịch vụ, Đơn vị tính, Số lượng, Đơn giá, Thành tiền. Các list con còn lại sẽ là các thông số về sản phẩm')
    ten_nguoi_ban: str = Field(default=None, description='Tên người bán hàng hoặc công ty bán hàng')
    ten_nguoi_mua: str = Field(default=None, description='Tên người mua hàng hoặc công ty mua hàng')
    ngay: Optional[str] = Field(default=None, description='Ngày lập hoá đơn')
    thang: Optional[str] = Field(default=None, description='Tháng lập hoá đơn')
    nam: Optional[str] = Field(default=None, description='Năm lập hoá đơn')
    ki_hieu: Optional[str] = Field(default=None, description='Ký hiệu hoá đơn')
    ma_so_thue_ban: Optional[str] = Field(default=None, description='Mã số thuế người bán')
    dia_chi_ban: Optional[str] = Field(default=None, description='Địa chỉ người bán')
    dien_thoai_ban: Optional[str] = Field(default=None, description='Điện thoại người bán')
    so_tai_khoan_ban: Optional[str] = Field(default=None, description='Số tài khoản người bán')
    ma_so_thue_mua: Optional[str] = Field(default=None, description='Mã số thuế người mua')
    dia_chi_mua: Optional[str] = Field(default=None, description='Địa chỉ người mua')
    dien_thoai_mua: Optional[str] = Field(default=None, description='Điện thoại người mua')
    so_tai_khoan_mua: Optional[str] = Field(default=None, description='Số tài khoản người mua')
    hinh_thuc_thanh_toan: Optional[str] = Field(default=None, description='Hình thức thanh toán')


def create_invoice_docx(invoice_data: InvoiceDetails, output_path: str) -> str:
    """
    Tạo một file DOCX hoá đơn từ đối tượng InvoiceDetails đã được trích xuất.
    # Node 'create_docx_node' sẽ gọi tool này.
    """
    print(f"--- TOOL: Đang tạo file DOCX tại: {output_path} ---")
    try:
        doc = Document()
        
        # Tiêu đề hóa đơn
        title_para = doc.add_paragraph()
        title_run = title_para.add_run('HÓA ĐƠN BÁN HÀNG TÍCH HỢP BIÊN LAI THU THUẾ, PHÍ, LỆ PHÍ')
        title_run.font.name = 'Times New Roman'
        title_run.font.size = 14
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thông tin thời gian và ký hiệu
        time_para = doc.add_paragraph()
        time_run = time_para.add_run(f'Ngày {invoice_data.ngay or "..."} tháng {invoice_data.thang or "..."} năm {invoice_data.nam or "..."}')
        time_run.font.name = 'Times New Roman'
        time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        symbol_para = doc.add_paragraph()
        symbol_run = symbol_para.add_run(f'Ký hiệu: {invoice_data.ki_hieu or "..."}')
        symbol_run.font.name = 'Times New Roman'
        symbol_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # # Phần I
        # section_para = doc.add_paragraph()
        # section_run = section_para.add_run('PHẦN I: HÓA ĐƠN BÁN HÀNG')
        # section_run.font.name = 'Times New Roman'
        # section_run.bold = True
        
        # Thông tin người bán
        doc.add_paragraph(f'Tên người bán: {invoice_data.ten_nguoi_ban}')
        if invoice_data.ma_so_thue_ban:
            doc.add_paragraph(f'Mã số thuế: {invoice_data.ma_so_thue_ban}')
        if invoice_data.dia_chi_ban:
            doc.add_paragraph(f'Địa chỉ: {invoice_data.dia_chi_ban}')
        if invoice_data.dien_thoai_ban:
            doc.add_paragraph(f'Điện thoại: {invoice_data.dien_thoai_ban}')
        if invoice_data.so_tai_khoan_ban:
            doc.add_paragraph(f'Số tài khoản: {invoice_data.so_tai_khoan_ban}')

        # Thông tin người mua
        doc.add_paragraph(f'Tên người mua: {invoice_data.ten_nguoi_mua}')
        if invoice_data.ma_so_thue_mua:
            doc.add_paragraph(f'Mã số thuế: {invoice_data.ma_so_thue_mua}')
        if invoice_data.dia_chi_mua:
            doc.add_paragraph(f'Địa chỉ: {invoice_data.dia_chi_mua}')
        if invoice_data.dien_thoai_mua:
            doc.add_paragraph(f'Điện thoại: {invoice_data.dien_thoai_mua}')
        if invoice_data.so_tai_khoan_mua:
            doc.add_paragraph(f'Số tài khoản: {invoice_data.so_tai_khoan_mua}')
        
        # Thông tin thanh toán
        if invoice_data.hinh_thuc_thanh_toan:
            doc.add_paragraph(f'Hình thức thanh toán: {invoice_data.hinh_thuc_thanh_toan}')
        doc.add_paragraph('Đồng tiền thanh toán: VNĐ')
        
        # Bảng hàng hóa
        if invoice_data.du_lieu_bang and len(invoice_data.du_lieu_bang) > 0:
            # Tạo bảng
            table = doc.add_table(rows=len(invoice_data.du_lieu_bang), cols=len(invoice_data.du_lieu_bang[0]))
            table.style = 'Table Grid'
            
            # Điền dữ liệu vào bảng
            for i, row in enumerate(invoice_data.du_lieu_bang):
                for j, value in enumerate(row):
                    cell = table.cell(i, j)
                    cell.text = str(value) if value is not None else ""
                    
                    # Định dạng header
                    if i == 0:  # Header row
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.name = 'Times New Roman'
        
        # Lưu file
        doc.save(output_path + '.docx')
        return f'Đã tạo thành công hóa đơn {output_path}.docx'
        
    except Exception as e:
        return f'Lỗi {e} khi tạo hóa đơn {output_path}'

def get_prompt_for_data_excel(data_excel_str: List[List[str]]) -> str:
    """Tạo prompt hướng dẫn LLM trích xuất thông tin từ dữ liệu Excel

    Args:
        data_excel_str (List[List[str]]): Dữ liệu Excel dưới dạng list of lists

    Returns:
        str: Prompt được tạo để hướng dẫn LLM
    """
    prompt = f"""
    DỮ LIỆU EXCEL CẦN PHÂN TÍCH:
    {data_excel_str}

    HÃY TRÍCH XUẤT THÔNG TIN SAU ĐÂY TỪ DỮ LIỆU EXCEL TRÊN:

    1. THÔNG TIN THỜI GIAN:
    - ngay: (ngày lập hóa đơn)
    - thang: (tháng lập hóa đơn) 
    - nam: (năm lập hóa đơn)
    - ki_hieu: (ký hiệu hóa đơn)

    2. THÔNG TIN NGƯỜI BÁN:
    - ten_nguoi_ban: (tên/công ty bán hàng)
    - ma_so_thue_ban: (mã số thuế người bán)
    - dia_chi_ban: (địa chỉ người bán)
    - dien_thoai_ban: (số điện thoại người bán)
    - so_tai_khoan_ban: (số tài khoản người bán)

    3. THÔNG TIN NGƯỜI MUA:
    - ten_nguoi_mua: (tên/công ty mua hàng)
    - ma_so_thue_mua: (mã số thuế người mua)
    - dia_chi_mua: (địa chỉ người mua)
    - dien_thoai_mua: (số điện thoại người mua)
    - so_tai_khoan_mua: (số tài khoản người mua)

    4. THÔNG TIN THANH TOÁN:
    - hinh_thuc_thanh_toan: (tiền mặt/chuyển khoản/...)

    5. THÔNG TIN HÀNG HÓA:
    - du_lieu_bang: Tạo bảng dữ liệu với cấu trúc [
        ['STT', 'Tên hàng hóa, dịch vụ', 'Đơn vị tính', 'Số lượng', 'Đơn giá', 'Thành tiền'],
        [dữ liệu hàng hóa từng dòng...]
    ]

    SAU KHI TRÍCH XUẤT, HÃY GỌI HÀM make_invoice VỚI CÁC THAM SỐ TRÍCH XUẤT ĐƯỢC.
    
    LƯU Ý: 
    - Nếu không tìm thấy thông tin nào, hãy để giá trị None hoặc chuỗi rỗng
    - Đảm bảo dữ liệu bảng hàng hóa có đúng định dạng list of lists
    - Tính toán lại thành tiền = số lượng x đơn giá nếu cần
    """
    
    return prompt

    
# make_invoice(file_name='hoa_don_test',
#              du_lieu_bang=[['STT', 'Tên hàng hoá, dịch vụ', 'Đơn vị tính', 'Số lượng', 'Đơn giá', 'Thành tiền'],
#                            ['1', 'Sản phẩm A', 'Cái', '2', '100000', '200000'],
#                            ['2', 'Sản phẩm B', 'Cái', '1', '150000', '150000'],
#                            ['3', 'Sản phẩm C', 'Cái', '3', '50000', '150000'],
#                            ['4', 'Sản phẩm D', 'Cái', '5', '20000', '100000']],
#              ngay='12',
#              thang='07',
#              nam='2023',
#              ki_hieu='HD001',
#              ten_nguoi_ban='Lê Việt Anh',
#              ma_so_thue_ban='0123456789',
#              dia_chi_ban='123 Đường ABC, Quận 1, TP.HCM',
#              dien_thoai_ban='0123456789',
#              so_tai_khoan_ban='12345678901234567890',
#              ten_nguoi_mua='Nguyễn Văn B',
#              ma_so_thue_mua='9876543210',
#              dia_chi_mua='456 Đường XYZ, Quận 2, TP.HCM',
#              dien_thoai_mua='0987654321',
#              so_tai_khoan_mua='09876543210987654321',
#              hinh_thuc_thanh_toan='Chuyển khoản ngân hàng')